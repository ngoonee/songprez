#!/usr/bin/env python

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import re
from songprez import spset


def _make_style(document, ptSize=12, fontName="Consolas", styleName="Custom Style"):
    textStyle = document.styles.add_style(styleName, WD_STYLE_TYPE.PARAGRAPH)
    textStyle.font.name = fontName
    textStyle.font.size = Pt(ptSize)
    textStyle.paragraph_format.space_before = 0
    textStyle.paragraph_format.space_after = 0
    textStyle.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return textStyle


def _append_songs(document, setpath, textStyle, lyricsOnly=False):
    mySet = spset.SPSet.read_from_file(setpath)
    songs = mySet.list_songs()
    for s in songs:
        p = document.add_paragraph(s.title, style=textStyle)
        p.runs[0].bold = True
        p.runs[0].underline = True
        document.add_paragraph("", style=textStyle)
        lyrics = s.words if lyricsOnly else s.lyrics
        document.add_paragraph(lyrics, style=textStyle)
        document.add_paragraph("", style=textStyle)
    return document

def write_docx(setpath, docpath, pianist=None, guitar=None, drum=None,
               tambourinists=None, colour=None):
    # Extract the worship leader's name by eliminating the initial date values
    basedir, filename = os.path.split(setpath)
    firstchar = re.search("[a-z,A-Z]", filename).start()
    worshipleader = filename[firstchar:]
    # Create a blank, default document and give it an appropriate title
    document = Document()
    document.core_properties.title = filename
    document.core_properties.comments = "Auto-generated song chords - songprez"
    # Modify orientation, margins, and column settings
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.left_margin = Inches(0.25)
    section.top_margin = Inches(0.25)
    section.right_margin = Inches(0.25)
    section.bottom_margin = Inches(0.25)
    sectPr = document.sections[0]._sectPr  # Get a pointer to the section
    sectPr.xpath('./w:cols')[0].set(qn('w:num'), "3")  # Set to 3 columns
    sectPr.xpath('./w:cols')[0].set(qn('w:space'), "360")  # Set to 0.25 inches
    # Create appropriate styles
    chordStyle = _make_style(document, 9, "Consolas", "Chords")
    lyricStyle = _make_style(document, 12, "Consolas", "Lyrics")
    document = _append_songs(document, setpath, chordStyle, False)
    p = document.add_paragraph("Worship Leader   ", style=chordStyle)
    p.add_run(worshipleader)
    p = document.add_paragraph("Piano            ", style=chordStyle)
    p.add_run(pianist)
    p = document.add_paragraph("Guitar           ", style=chordStyle)
    p.add_run(guitar)
    p = document.add_paragraph("Drum             ", style=chordStyle)
    p.add_run(drum)
    p = document.add_paragraph("Tambourinists    ", style=chordStyle)
    p.add_run(tambourinists)
    p = document.add_paragraph("Colour           ", style=chordStyle)
    p.add_run(colour)
    p = document.add_page_break()
    document = _append_songs(document, setpath, lyricStyle, True)
    document.save(docpath)
